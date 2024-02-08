from turtle import position
from docx import Document
from docx.shared import Pt, RGBColor
import win32com.client as win32
import pathlib
import configparser
import requests
from bs4 import BeautifulSoup
import regex as re
import subprocess
import time
from collections import Counter
import random
import sys

# Read the settings values value from the settings file
config = configparser.ConfigParser()
config.read('C:\\Users\\Maryl\\Downloads\\dice-easy-apply\\chrome_driver\\settings.ini')
resume_pdf = "" + config.get('mode', 'resume_pdf') + ""
resume_folder = config.get('files', 'resume_location')
directory = config.get('home', 'working_directory')
resume_name = config.get('files', 'resume_name')
original_resume = config.get('files', 'original_resume')
search_skills_jd = config.get('files', 'Search_Phrase_Resume')

with open(f"{directory}url.txt", "r") as file:
    base_url = file.read()        
    file.close()

#base_url = str(sys.argv[1])
#print(base_url)
base_url = None


def job_description(url):
    response = requests.get(url)  # Send a GET request to the URL
    soup = BeautifulSoup(response.text, 'html.parser')  # Parse the HTML content

    links = []
    for link in soup.find_all('div', {'data-testid': 'jobDescriptionHtml'}):  # Find all <div> tags with data-testid 'jobDescriptionHtml'
        text = link.text  # Get the display text of the link
        links.append(text)  # Add the text to the links list

    with open(f"{directory}dicejobdescription.txt", 'w') as file:
        file.write(str(links))


def skills(url):
    response = requests.get(url)  # Send a GET request to the URL
    soup = BeautifulSoup(response.text, 'html.parser')  # Parse the HTML content
    
    links = []
    for link in soup.find_all('span', id=lambda value: value and value.startswith('skillChip')):  # Find all <span> tags with id starting with 'skillChip'
        text = link.text  # Get the display text of the link
        links.append(text)  # Add the text to the links list
    
    with open(f"{directory}skillslist.txt", 'w') as file:
        file.writelines('\n'.join(links))

def compare_jobdescription():
    words1 = set()  # Set to store unique words before ":"

    # Open the file in read mode
    with open(f"{directory}skillslist.log", 'r') as file:
        # Read each line
        for line in file:
            # Find the position of the first occurrence of ":"
            delimiter_index = line.find(':')
            # If ":" is found, append the part before it to the set
            if delimiter_index != -1:
                words1.add(line[:delimiter_index])

    matching_words_jobdescription = [] # Set to store unique matching words

    with open(f"{directory}dicejobdescription.txt", "r") as file:
        for line in file:
            # Split the line into words
            #words2 = line.split()
            for word in words1:
                # Compare each word in file2 with words in file1
                if word in line:
                    #matching_words_jobdescription.add(word)
                    matching_words_jobdescription.append(word)
                    #print(matching_words_jobdescription)


    matching_words_skills = []  # Set to store unique matching words

    with open(f"{directory}skillslist.txt", "r") as file:
        for line in file:
            # Split the line into words
            #words2 = line.split()
            for word in words1:
                # Compare each word in file2 with words in file1
                if word in line:
                    #matching_words_skills.add(word)
                    matching_words_skills.append(word)
                    #print(matching_words_skills)

#    set1 = {1, 2, 3}
#    set2 = {3, 4, 5}
    
    combined_list = matching_words_jobdescription + matching_words_skills

    #print("combined:", combined_list)

    word_counts = Counter(combined_list)
    sorted_words = sorted(word_counts, key=lambda w: word_counts[w], reverse=True)
    #print(sorted_words)

    #resume_lines = set()
    new_skills = []
    #print(matching_words)
    with open(f"{directory}skillslist.log", 'r') as file:
        for line in file:
            #word3 = line.split()
            delimiter_index = line.find(':')
            
            #print(line)
            if delimiter_index != -1:

                skills = [line[delimiter_index + 1:].strip() for match in sorted_words if line.startswith(f"{match}:")]

                new_skills.extend(element for element in skills if element)

    list_length = len(new_skills)
    skills_needed = int(20)
    #print(list_length, skills_needed)
    if list_length < skills_needed:
        skills_list = new_skills
        get_skills = (skills_needed - list_length)
        #print(get_skills)
        random_lines = []
        #print(matching_words)
        with open(f"{directory}skillslist.log", 'r') as file:
            lines = file.readlines()
            #print(lines)
            for line in lines:
                word3 = line.split()
                delimiter_index = line.find(':')
                
                #print(line)
                if delimiter_index != -1:

                    random_lines = random.sample(lines[:128], get_skills) 
                    counter = 0
                    for line in random_lines:
                        counter += 1    
                        for match in skills_list:
                            if not line.startswith(f"{match}:"):
                                        random_line = line.split(':', maxsplit=1)[-1].strip()
                                        random_lines.extend(element for element in random_line if element)
                                        #print(random_lines)
                                        #print(get_skills)
                        if counter == get_skills:
                                        break
                                        
                        break
                    break
                
                                        
             
                        #random_lines.extend(element for element in random_lines if element)
        #random_lines = line.split(':', maxsplit=1)[-1].strip()
        del random_lines[get_skills:]
        random_lines = [item.replace("\n", "").strip() for item in random_lines]
        random_lines = [item.split(":")[1].strip() for item in random_lines]
        #print("random_skills:", random_lines)
        added_skills = []            
        join = skills_list + random_lines

        added_skills.extend(element for element in join if element)
        #new_skills.append(join)
        new_skills = added_skills
        #print(new_skills)
        #count = len(new_skills)
        #print(count)
    
    return new_skills


def search_replace(doc, search_term, replace_term, font_size):
    #print(search_term, "1")
    for paragraph in doc.paragraphs:
        if "small" in font_size:  
            if search_term in paragraph.text:
                #print(search_term, "3")
                inline = paragraph.runs
                for i in range(len(inline)):
                    if search_term in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_term, replace_term)
                        font = inline[i].font   
                        font.size = Pt(1)  # Set font size to 1 point
                        font.color.rgb = RGBColor(255, 255, 255)  # Set color to white
                        
        if "big" in font_size:
            if str(search_term) in paragraph.text:
                #print(search_term, "2")
                inline = paragraph.runs
                for i in range(len(inline)):
                    if search_term in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_term, enumerate(replace_term))
                        font = inline[i].font
                        # Set font to Garamond and size 11
                        font.size = Pt(11)  # Set font size to 11 point
                        #font.name = "Garamond"  # Set font style to Garamond

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if str(search_term) in cell.text:
                    #print(search_term, "4")
                    cell.text = cell.text.replace(search_term, enumerate(replace_term))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            if "small" in font_size:
                                font.size = Pt(1)  # Set font size to 1 point
                                font.color.rgb = RGBColor(255, 255, 255)  # Set color to white
                            if "big" in font_size:
                                font.size = Pt(11)  # Set font size to 11 point
                                font.name = "Garamond"  # Set font style to Garamond

def search_skills(doc, search_terms, replace_terms):
    #print(search_terms)
    #print(replace_terms)\

    for term_index, description in enumerate(search_terms):
        for paragraph in doc.paragraphs:
            if search_terms[term_index] in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if search_terms[term_index] in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_terms[term_index], replace_terms[term_index])


        # Search and replace in tables
     
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_terms[term_index] in cell.text:
                        #print(search_terms)
                        for paragraph in cell.paragraphs:
                            #for run in paragraph.runs:
                            if search_terms[term_index] in paragraph.text:  # Check if run contains the search term
                                paragraph.text = paragraph.text.replace(search_terms[term_index], replace_terms[term_index])
                                for run in paragraph.runs:
                                    run.font.size = Pt(11)  # Modify font size to 11 point
                                    run.font.name = "Garamond"  # Modify font style to Garamond

def remove_unused_skills(doc, search_terms):

    #for term_index, description in enumerate(search_terms):
        for table in doc.tables:
            # Iterate through each row in the table
            for row in table.rows:
                # Iterate through each cell in the row
                for cell in row.cells:
                    #if search_terms[term_index] in cell.text:
                        # Iterate through each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        # Iterate through each line in the paragraph
                        for line in paragraph.text.splitlines():
                            # Check if the line contains any of the search terms
                            if any(term in line for term in search_terms):
                                # Remove the line by replacing it with an empty string
                                paragraph.text = paragraph.text.replace(line, '')

    #doc.save(file_path)
    


def create_search_terms():
    search_terms = []
    #search_terms2 = []
    for i in range(0, 10):
        search_terms.append(f" First{i}")
    
    for i in range(0, 10):
        search_terms.append(f" Second{i}")
    #search_terms = [str(item) for item in search_terms]
    #print(search_terms)


    #print(search_terms2)
    return search_terms

#job_description(base_url)
#skills(base_url)
## Open the Word document
##doc = Document(f"{resume_folder}{original_resume}.docx")
doc = Document(f"{resume_folder}{original_resume}2.docx")
#
#with open(f"{directory}skillslist.txt", "r") as file:
#    skills_list = file.read()
#
#with open(f"{directory}dicejobdescription.txt", "r") as file:
#    dicejobdescription = file.read()
#
#skills_list_and_jobdescription = str((skills_list, dicejobdescription))
#search_skill_description = str(f"{search_skills_jd}")
##print(skills_list_and_jobdescription)
#font_size = "small"
## Perform search and replace
#search_replace(doc, search_skill_description, skills_list_and_jobdescription, font_size)


matching_jd = compare_jobdescription()
#print(matching_jd)
replace_list = matching_jd
replace_number = len(replace_list)
search_list = create_search_terms()
#for i in range(0, 20):
#    print(replace_list[i])
replace_number = len(replace_list)

if replace_number > 20:
    
    del replace_list[20:replace_number] 
    replace_number = len(replace_list)

#replace_number = (len(replace_number))

if replace_number != 20:
    skills_range_search = search_list[replace_number:20]
    #print(skills_range_search)
    remove_unused_skills(doc, skills_range_search)


search_skills(doc, search_list, replace_list)

# Save the modified document
doc.save(f"{resume_folder}{resume_name}.docx")
#time.sleep(1)

# Initialize the Word application
word_app = win32.gencache.EnsureDispatch('Word.Application')

# Open the Word document via the application
word_doc = word_app.Documents.Open(f"{resume_folder}{resume_name}.docx")

if "yes" in resume_pdf:
    # Save the Word document as PDF
    word_doc.SaveAs(f"{resume_folder}{resume_name}.pdf", FileFormat=17)

    # Close the Word document and application
    word_doc.Close()
    word_app.Quit()

#    del_path = pathlib.Path(f'{resume_folder}{resume_name}.docx')
#    del_path.unlink()