from turtle import position
from docx import Document
from docx.shared import Pt, RGBColor
import win32com.client as win32
import pathlib
import configparser
import requests
from bs4 import BeautifulSoup
import regex as re
import subprocess
import sys

# Read the settings values value from the settings file
config = configparser.ConfigParser()
config.read('C:\\Users\\Maryl\\Downloads\\dice-easy-apply\\chrome_driver\\settings.ini')
coverletter_pdf = "" + config.get('mode', 'coverletter_pdf') + ""
directory = config.get('home', 'working_directory')
coverletter_name = config.get('files', 'coverletter_name')
coverletter_folder = config.get('files', 'coverletter_location')
search_term_company = config.get('files', 'Search_term_Company')
search_term_position = config.get('files', 'Search_term_Position')
original_coverletter = config.get('files', 'original_coverletter')
search_term_skills = config.get('files', 'Search_term_Skills')
search_term_manager = config.get('files', 'Search_term_Manager')

with open("url.txt", "r") as file:
    base_url = file.read()        
    file.close()

#base_url = str(sys.argv[1])
#print(base_url)

def job_info(url):
    response = requests.get(url)  # Send a GET request to the URL
    soup = BeautifulSoup(response.text, 'html.parser')  # Parse the HTML content

    elements = [('h1', {'data-cy': 'jobTitle'}), ('a', {'data-cy': 'companyNameLink'})]
    links = []
    for element in elements:
        for link in soup.find_all(element[0], element[1]):
            links.append(link.text)
    with open(f"{directory}coverletter-info.txt", 'w') as file:
        file.writelines('\n'.join(links))
        


def skills(url):
    response = requests.get(url)  # Send a GET request to the URL
    soup = BeautifulSoup(response.text, 'html.parser')  # Parse the HTML content
    
    links = []
    for link in soup.find_all('span', id=lambda value: value and value.startswith('skillChip')):  # Find all <span> tags with id starting with 'skillChip'
        text = link.text  # Get the display text of the link
        links.append(text)  # Add the text to the links list
    
    with open(f"{directory}skillslist.txt", 'w') as file:
        file.writelines('\n'.join(links))



def email(url):
    response = requests.get(url)  # Send a GET request to the URL
    soup = BeautifulSoup(response.text, 'html.parser')  # Parse the HTML content

    links = []
    for link in soup.find_all('script', {'id': '__NEXT_DATA__'}):  # Find all <div> tags with data-testid 'jobDescriptionHtml'
        text = link.text  # Get the display text of the link
        links.append(text)  # Add the text to the links list

    content = str(links)  # Replace with the actual content you want to extract the email from

    email_match = re.search(r'(?<=applicationDetail":{"email":")(.*?)(?="ccEmail")', content)
    if email_match:
        email = email_match.group(1)
        email_prefix = email.split('@')[0]

        if email_prefix:
            email_prefix = email_prefix[0].upper() + email_prefix[1:]

            if '.' in email_prefix:
                email_prefix = email_prefix.replace('.', ' ')
                email_prefix = ' '.join(word.capitalize() for word in email_prefix.split(' '))

            if '_' in email_prefix:
                email_prefix = ''

        if email_prefix != '':
            with open(f"{directory}coverletter-info.txt", 'a') as file:
                file.writelines('\n' + email_prefix )



def search_replace(doc, search_terms, replace_terms):
    for term_index, description in enumerate(search_terms):
        for paragraph in doc.paragraphs:
            if search_terms[term_index] in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if search_terms[term_index] in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_terms[term_index], replace_terms[term_index])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_terms[term_index] in cell.text:
                        cell.text = cell.text.replace(search_terms[term_index], replace_terms[term_index])

job_info(base_url)
skills(base_url)
email(base_url)
# Open the Word document
doc = Document(f"{coverletter_folder}{original_coverletter}.docx")

with open(f"{directory}skillslist.txt", "r") as skills:
    skills_list = skills.readlines()
    skills_range = ", ".join([skill.strip() for skill in skills_list])
    #print(skills_range)


    with open(f"{directory}coverletter-info.txt", "r") as file:
        position = file.readline()
        company = file.readline()
        hiring_manager = file.readline()
        if hiring_manager:
            search_terms = [f"{search_term_position}", f"{search_term_company}", f"{search_term_skills}", f"{search_term_manager}"]
            replace_terms = [position.strip(), company.strip(), skills_range.strip(), hiring_manager.strip()]
        else:
            search_terms = [f"{search_term_position}", f"{search_term_company}", f"{search_term_skills}"]
            replace_terms = [position.strip(), company.strip(), skills_range.strip()]

# Perform search and replace
search_replace(doc, search_terms, replace_terms)


# Save the modified document
doc.save(f"{coverletter_folder}{coverletter_name}.docx")
#time.sleep(1)

# Initialize the Word application
word_app = win32.gencache.EnsureDispatch('Word.Application')

# Open the Word document via the application
word_doc = word_app.Documents.Open(f"{coverletter_folder}{coverletter_name}.docx")

if "yes" in coverletter_pdf:
    # Save the Word document as PDF
    word_doc.SaveAs(f"{coverletter_folder}{coverletter_name}.pdf", FileFormat=17)

    # Close the Word document and application
    word_doc.Close()
    word_app.Quit()

    del_path = pathlib.Path(f'{coverletter_folder}{coverletter_name}.docx')
    del_path.unlink()